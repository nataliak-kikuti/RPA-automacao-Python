from docx import Document

#Aunmeta o tamanho da letra
from docx.shared import Pt

from openpyxl import load_workbook


#Muda a cor do texto
from docx.shared import RGBColor

import os

import win32com.client as win32
outlook = win32.Dispatch("outlook.application")



nome_arquivo_alunos ="DadosAlunosEmail.xlsx"
planilhaDadosAlunos = load_workbook(nome_arquivo_alunos)

sheet_selecionada = planilhaDadosAlunos["Nomes"]

for linha in range(2, len(sheet_selecionada["A"]) +1):
    arquivoWord = Document("D:\\Usuários\\Natalia\\Desktop\\RPA\\Word\\Certificado3.docx")

    estilo = arquivoWord.styles["Normal"]

    nomeAluno = sheet_selecionada["A%s" % linha].value
    dia = sheet_selecionada["B%s" % linha].value
    mes = sheet_selecionada["C%s" % linha].value
    ano = sheet_selecionada["D%s" % linha].value
    nomeCurso = sheet_selecionada["E%s" % linha].value
    nomeInstrutor = sheet_selecionada["F%s" % linha].value
    emailAluno = sheet_selecionada["G%s" % linha].value


    for paragrafo in arquivoWord.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        paragrafoP1 = "Concluiu com sucesso o curso de "
        paragrafoP2 = ", como carga horária de 20 horas, promovido pela escola de Cursos Online em  "
        terceiraParteParagrafo = f" {paragrafoP2} {dia} de {mes} de {ano}."

        if "escola" in paragrafo.text:
            paragrafo.text = paragrafoP1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicionaNovaPalavra = paragrafo.add_run(nomeCurso)
            adicionaNovaPalavra.font.color.rgb = RGBColor(255, 0, 0)
            adicionaNovaPalavra.underline = True
            adicionaNovaPalavra.bold = True
            adicionaNovaPalavra = paragrafo.add_run(terceiraParteParagrafo)
            adicionaNovaPalavra.font.color.rgb = RGBColor(0, 0, 0)

        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeInstrutor + "- Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminhoCertificados = "D:\\Usuários\\Natalia\\Desktop\\RPA\\Word\\Certificados" + nomeAluno + ".docx"

    # Salvando o certificado do aluno
    arquivoWord.save(caminhoCertificados)
    emailOutlook = outlook.CreateItem(0)
    emailOutlook.To = emailAluno
    emailOutlook.Subject = "Certificado " + nomeAluno
    emailOutlook.HtmlBody = f"""
    <p>Boa noite {nomeAluno}. </p>
    <p>Segue seu <b>certificado.</b> </p>
    <p>Atenciosamente,</p>
"""
    emailOutlook.Attachments.Add(caminhoCertificados)
    emailOutlook.save()

print("Certificados gerados com sucesso!")


